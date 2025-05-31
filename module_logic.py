"""
Logic for the Input Processing & Ingestion Module.

This module is responsible for:
- Robustly handling various document formats (PDF, DOCX, TXT).
- Implementing OCR for scanned documents or images within documents.
- Performing text cleaning and basic pre-processing.
- Leveraging Manus capabilities and free tools.
"""

import os
import subprocess
import re
from PIL import Image # For image processing with Tesseract
import pytesseract # For OCR
import docx # For .docx files

# For PDF processing, we'll use pdftotext (from poppler-utils, assumed installed)
# and Tesseract for image-based PDFs or images within PDFs.

class InputProcessor:
    def __init__(self):
        """Initializes the InputProcessor."""
        print("InputProcessor initialized.")
        # Ensure tesseract is in PATH or specify its location if needed
        # Example: pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'

    def _clean_text(self, text):
        """Basic text cleaning: remove excessive newlines, leading/trailing whitespace."""
        if not text:
            return ""
        text = re.sub(r"\n{3,}", "\n\n", text) # Replace 3+ newlines with 2
        text = text.strip()
        return text

    def process_txt(self, file_path):
        """Processes a .txt file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return self._clean_text(text)
        except Exception as e:
            print(f"Error processing TXT file {file_path}: {e}")
            return None

    def process_docx(self, file_path):
        """Processes a .docx file."""
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return self._clean_text("\n".join(full_text))
        except Exception as e:
            print(f"Error processing DOCX file {file_path}: {e}")
            return None

    def _ocr_image_to_text(self, image_path):
        """Performs OCR on a single image file."""
        try:
            text = pytesseract.image_to_string(Image.open(image_path))
            return self._clean_text(text)
        except Exception as e:
            print(f"Error performing OCR on image {image_path}: {e}")
            return ""

    def process_pdf(self, file_path):
        """
        Processes a .pdf file.
        First attempts to extract text directly using pdftotext.
        If that yields little or no text (suggesting an image-based PDF),
        it then attempts to convert PDF pages to images and OCR them.
        """
        extracted_text = ""
        try:
            # Attempt direct text extraction using pdftotext
            result = subprocess.run(["pdftotext", file_path, "-"], capture_output=True, text=True, check=True)
            extracted_text = self._clean_text(result.stdout)
        except subprocess.CalledProcessError as e:
            print(f"pdftotext error for {file_path}: {e}. Will attempt OCR.")
        except FileNotFoundError:
            print("pdftotext command not found. Please ensure poppler-utils is installed. Will attempt OCR if possible.")
        except Exception as e:
            print(f"Error during pdftotext processing for {file_path}: {e}. Will attempt OCR.")

        # If direct text extraction is insufficient (e.g., less than 100 characters for a multi-page doc, or failed),
        # or if we want to supplement with OCR for images within the PDF.
        # For simplicity here, if pdftotext fails or gives very short text, we try full OCR.
        # A more advanced version would detect images within text-based PDFs and OCR them selectively.
        if not extracted_text or len(extracted_text) < 100: # Arbitrary threshold for "insufficient text"
            print(f"Direct text extraction from {file_path} was insufficient or failed. Attempting OCR...")
            ocr_texts = []
            temp_image_dir = "/tmp/pdf_pages_for_ocr"
            os.makedirs(temp_image_dir, exist_ok=True)
            
            try:
                # Convert PDF pages to images (e.g., PPM using pdftoppm from poppler-utils)
                subprocess.run(["pdftoppm", "-png", file_path, os.path.join(temp_image_dir, "page")], check=True)
                image_files = sorted([os.path.join(temp_image_dir, f) for f in os.listdir(temp_image_dir) if f.startswith("page") and f.endswith(".png")])
                
                if not image_files:
                    print(f"No images generated from PDF {file_path} for OCR.")
                    # Fallback to any text extracted by pdftotext earlier, even if minimal
                    return extracted_text if extracted_text else None 

                for i, img_path in enumerate(image_files):
                    print(f"Performing OCR on page {i+1} of {file_path}...")
                    page_text = self._ocr_image_to_text(img_path)
                    ocr_texts.append(page_text)
                    os.remove(img_path) # Clean up temporary image file
                
                # Clean up directory if empty
                if not os.listdir(temp_image_dir):
                    os.rmdir(temp_image_dir)
                
                extracted_text = "\n\n--- Page Break ---\n\n".join(ocr_texts) # Join OCRed text from all pages
                extracted_text = self._clean_text(extracted_text)
            except subprocess.CalledProcessError as e:
                print(f"pdftoppm error for {file_path}: {e}. OCR attempt failed.")
                return extracted_text if extracted_text else None # Return whatever pdftotext got, or None
            except FileNotFoundError:
                print("pdftoppm command not found. Please ensure poppler-utils is installed. OCR attempt failed.")
                return extracted_text if extracted_text else None
            except Exception as e:
                print(f"Error during PDF to image conversion or OCR for {file_path}: {e}")
                return extracted_text if extracted_text else None
        
        return extracted_text if extracted_text else None

    def process_document(self, file_path):
        """Detects file type and processes accordingly."""
        if not os.path.exists(file_path):
            print(f"Error: File not found at {file_path}")
            return None

        _, extension = os.path.splitext(file_path.lower())
        print(f"Processing document: {file_path} (type: {extension})")

        if extension == ".txt":
            return self.process_txt(file_path)
        elif extension == ".docx":
            return self.process_docx(file_path)
        elif extension == ".pdf":
            return self.process_pdf(file_path)
        elif extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"]:
            print("Processing as a direct image file for OCR.")
            return self._ocr_image_to_text(file_path)
        else:
            print(f"Unsupported file type: {extension} for file {file_path}")
            return None

# Example usage (for testing purposes within this file)
if __name__ == "__main__":
    processor = InputProcessor()
    test_docs_dir = "/home/ubuntu/ai_rfp_agent/test_documents"
    os.makedirs(test_docs_dir, exist_ok=True)

    # Create a sample TXT file
    sample_txt_path = os.path.join(test_docs_dir, "sample.txt")
    with open(sample_txt_path, "w") as f:
        f.write("This is a sample text file.\nIt has multiple lines.\n\nAnd some extra spaces.   ")
    
    # Create a sample DOCX file (requires python-docx to be installed)
    try:
        from docx import Document
        sample_docx_path = os.path.join(test_docs_dir, "sample.docx")
        doc = Document()
        doc.add_paragraph("This is a sample DOCX file content.")
        doc.add_paragraph("It also has paragraphs and could include tables or images (not tested here).")
        doc.save(sample_docx_path)
    except ImportError:
        print("python-docx library not found, skipping DOCX test file creation.")
        sample_docx_path = None

    # For PDF and Image OCR, manual test files would be needed in test_docs_dir.
    # e.g., a text-based PDF, an image-based PDF, and a standalone PNG image.
    # Create a placeholder for a PDF to demonstrate the call, assuming one exists.
    # To test PDF OCR, you would need to place a PDF (e.g., image-based) at sample_pdf_path
    sample_pdf_path = os.path.join(test_docs_dir, "sample_image.pdf") # Needs to be an actual image PDF for OCR test
    # Create a dummy PDF if poppler-utils and a real PDF are not available for basic path testing
    if not os.path.exists(sample_pdf_path):
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(sample_pdf_path)
            c.drawString(100, 750, "This is a minimal PDF for path testing.")
            c.save()
            print(f"Created dummy PDF for testing: {sample_pdf_path}")
        except ImportError:
            print("reportlab not found, cannot create dummy PDF. PDF processing might fail if no sample_image.pdf exists.")

    print("\n--- Testing Input Processor ---")
    
    txt_content = processor.process_document(sample_txt_path)
    print(f"\nContent from TXT ({sample_txt_path}):\n---\n{txt_content}\n---")

    if sample_docx_path and os.path.exists(sample_docx_path):
        docx_content = processor.process_document(sample_docx_path)
        print(f"\nContent from DOCX ({sample_docx_path}):\n---\n{docx_content}\n---")
    
    # Test PDF (will attempt pdftotext first, then OCR if needed and sample_image.pdf is image-based)
    if os.path.exists(sample_pdf_path):
        pdf_content = processor.process_document(sample_pdf_path)
        print(f"\nContent from PDF ({sample_pdf_path}):\n---\n{pdf_content}\n---")
    else:
        print(f"\nSkipping PDF test as {sample_pdf_path} does not exist.")

    # Example of direct image OCR (place an image file here to test)
    # sample_image_path = os.path.join(test_docs_dir, "sample_ocr_image.png") 
    # if os.path.exists(sample_image_path):
    #     image_ocr_content = processor.process_document(sample_image_path)
    #     print(f"\nOCR Content from Image ({sample_image_path}):\n---\n{image_ocr_content}\n---")

    print("\nInput Processing & Ingestion module logic updated with free tools.")

